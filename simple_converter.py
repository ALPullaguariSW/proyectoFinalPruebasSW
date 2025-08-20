#!/usr/bin/env python3
"""
Convertidor simple de Markdown a Word
Requerimientos: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os

def create_word_document(md_file, output_file):
    """Crear documento Word desde Markdown"""
    
    # Leer el archivo Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Crear documento Word
    doc = Document()
    
    # Configurar estilos
    styles = doc.styles
    
    # Estilo para títulos principales
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    # Estilo para subtítulos
    subtitle_style = styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = 'Arial'
    subtitle_style.font.size = Pt(16)
    subtitle_style.font.bold = True
    
    # Estilo para encabezados
    header_style = styles.add_style('CustomHeader', WD_STYLE_TYPE.PARAGRAPH)
    header_style.font.name = 'Arial'
    header_style.font.size = Pt(14)
    header_style.font.bold = True
    
    # Estilo para código
    code_style = styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)
    
    # Procesar contenido línea por línea
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if line.startswith('# '):
            # Título principal
            title = line[2:]
            p = doc.add_paragraph(title, style='CustomTitle')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif line.startswith('## '):
            # Subtítulo
            subtitle = line[3:]
            p = doc.add_paragraph(subtitle, style='CustomSubtitle')
            
        elif line.startswith('### '):
            # Encabezado
            header = line[4:]
            p = doc.add_paragraph(header, style='CustomHeader')
            
        elif line.startswith('|'):
            # Tabla
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) >= 2:
                # Crear tabla
                rows = len(table_lines)
                cols = len(table_lines[0].split('|')) - 2
                
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                for row_idx, table_line in enumerate(table_lines):
                    cells = table_line.split('|')[1:-1]
                    for col_idx, cell_content in enumerate(cells):
                        if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                            table.rows[row_idx].cells[col_idx].text = cell_content.strip()
            
            continue
            
        elif line.startswith('```'):
            # Bloque de código
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text, style='CustomCode')
                p.paragraph_format.left_indent = Inches(0.5)
            
        elif line.startswith('- ') or line.startswith('* '):
            # Lista
            list_item = line[2:]
            p = doc.add_paragraph(list_item, style='List Bullet')
            
        elif line.startswith('1. '):
            # Lista numerada
            list_item = line[3:]
            p = doc.add_paragraph(list_item, style='List Number')
            
        elif line == '---':
            # Separador
            doc.add_paragraph('_' * 50)
            
        elif line.startswith('**') and line.endswith('**'):
            # Texto en negrita
            bold_text = line[2:-2]
            p = doc.add_paragraph(bold_text)
            p.runs[0].font.bold = True
            
        elif line.startswith('✅') or line.startswith('❌'):
            # Emojis con texto
            text = line[1:].strip()
            p = doc.add_paragraph(text)
            
        elif line and not line.startswith('[') and not line.startswith('!['):
            # Párrafo normal
            if line:
                p = doc.add_paragraph(line)
        
        i += 1
    
    # Guardar documento
    doc.save(output_file)
    print(f"✅ Documento Word creado: {output_file}")

def main():
    """Función principal"""
    
    md_file = "INFORME_TECNICO_PROYECTO_3P.md"
    word_file = "INFORME_TECNICO_PROYECTO_3P.docx"
    
    if not os.path.exists(md_file):
        print(f"❌ Error: No se encontró el archivo {md_file}")
        return
    
    print("🚀 Iniciando conversión del informe técnico a Word...")
    
    try:
        create_word_document(md_file, word_file)
        print(f"\n🎉 Conversión completada!")
        print(f"📄 Word: {word_file}")
        print("\n💡 Para crear PDF:")
        print("   1. Abre el archivo Word")
        print("   2. Archivo → Exportar → Crear PDF")
        print("   3. O usa: File → Save As → PDF")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        print("\n💡 Instala las dependencias:")
        print("   pip install python-docx")

if __name__ == "__main__":
    main()
